"""
JEPA-PREDICT 预训练增强方案 — 6个思路 + 参考论文
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

# ── Title ──
title=doc.add_heading('JEPA 预训练增强方案\n— 基于8篇最新论文 (2024-2025)',level=0)
title.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('2026-06-26 | 与当前 JEPA+M2AE+JETS 架构兼容').alignment=WD_ALIGN_PARAGRAPH.CENTER

# ══════ 1. Token Level Alignment ══════
doc.add_heading('思路一：Token级跨模态对齐 (替换 InfoNCE)', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'Shi et al., "Are Electrodes Universal? Resolution-Agnostic Tokenization'
    ' for Time Series Cross-Modal Foundation Models," arXiv 2025.\n\n'
    '核心观点: 不同模态(ECG/PPG)的时序token在相同时间位置存在 '
    '"生理对应关系"——ECG的R峰token应对齐PPG的收缩期token。\n'
    '逐token对齐比全局pooled对齐细粒度62倍。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    '当前: InfoNCE(pooled_ECG, pooled_PPG) — 全局对比, 需要大batch\n\n'
    '改进:\n'
    '  1. 两个编码器输出的token序列 (B, 62, 512)\n'
    '  2. 对每个token位置i, 计算 cosine_sim(ecg_token_i, ppg_token_i)\n'
    '  3. L_align = mean(1 - cosine_sim) / 62\n'
    '  4. 替代InfoNCE, 无需负样本, batch size任意\n\n'
    '改动量: ~30行 (替换_conput_contrast_loss)'
)

doc.add_heading('预期收益', level=2)
doc.add_paragraph(
    '不受batch size限制; 学到每个时间位置的ECG-PPG形态对应;'
    ' CHD特征(如P波延迟→PPG收缩延迟)在token级别更明显'
)

# ══════ 2. Curriculum Masking ══════
doc.add_heading('思路二：Curriculum Masking (递增掩码难度)', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'Bengio et al., "Curriculum Learning," ICML 2009 (经典).\n'
    'Shi et al., "Are Electrodes Universal?," 2025 \n'
    '  — 提出multi-granularity masking: 粗→细粒度递增\n\n'
    '核心观点: 从简单任务开始, 逐步增加难度, 比固定难度收敛更快、泛化更好。'
    '生理信号天然适合——先学整体节律, 再学局部形态。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    '当前: 固定70%随机patch掩码\n\n'
    '改进:\n'
    '  0-15 epoch:  40% 掩码 + 块状掩码(连续一大块, 更容易)\n'
    '  16-30 epoch: 55% 掩码 + 混合掩码\n'
    '  31-50 epoch: 70% 掩码 + JETS随机patch\n'
    '  51-100 epoch: 85% 掩码 + JETS随机patch\n\n'
    '改动量: ~10行 (在train_pretrain.py中按epoch调整mask_ratio)'
)

doc.add_heading('预期收益', level=2)
doc.add_paragraph(
    '更快的初始收敛; 更强最终泛化; 零风险、低改动'
)

# ══════ 3. Frequency Masking ══════
doc.add_heading('思路三：频域掩码 (双域互补)', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'WavesFM (ICLR 2025) — MIT+Harvard, PPG foundation model\n'
    '  — 同时做时域MAE和频域MAE, 双域互补\n'
    '  — 证明频域特征对HRV、血压等心血管任务至关重要\n\n'
    '核心观点: 时域掩码学波形形态, 频域掩码学功率谱结构。'
    '两者互补——CHD患者的频域特征(HRV低频/高频比)本身就是诊断金标准。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    '新增频域重建辅助任务:\n'
    '  1. 对ECG做FFT → 得到频谱 (B, 1, L/2)\n'
    '  2. 随机遮掉某段频带(如0.04-0.15Hz, 对应LF)\n'
    '  3. Decoder重建完整频谱 → MSE(预测频谱, 真实频谱)\n'
    '  4. L = L_JEPA + L_InfoNCE + 0.1×L_freq\n\n'
    '改动量: ~40行 (加FrequencyHead + 频域损失)'
)

doc.add_heading('预期收益', level=2)
doc.add_paragraph('直接学习CHD金标准特征(HRV频域); 编码器更全面')

# ══════ 4. Two-Stage Pretraining ══════
doc.add_heading('思路四：两阶段预训练 (MAE → JEPA)', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'ECGFounder-PT (2025) — Li et al.\n'
    '  — 提出post-training策略: Stage1=Linear Probe → Stage2=Regularized FT\n'
    '  — 两阶段显著优于单阶段 (+1.2%~+3.3% macro AUROC)\n\n'
    '核心思想应用于预训练:\n'
    '  Stage 1 (重建): 先学"信号长什么样"\n'
    '  Stage 2 (预测): 再学"信号之间的因果"\n'
    '类比: 先学会认字母, 再学造句。比直接学造句更高效。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    'Stage 1 (epoch 0-20): 纯MAE重建\n'
    '  L = MSE(decoder(encoder(masked_ecg)), original_ecg)\n'
    '       + MSE(decoder(encoder(masked_ppg)), original_ppg)\n'
    '  → 编码器学基本的波形形态\n\n'
    'Stage 2 (epoch 21-100): JEPA预测 + InfoNCE (当前方案)\n'
    '  → 在形态理解基础上做跨模态推理\n\n'
    '改动量: ~20行 (加一个decoder, 前20 epoch用MAE loss)'
)

# ══════ 5. CPC ══════
doc.add_heading('思路五：时序对比预测 (CPC辅助任务)', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'Oord et al., "Representation Learning with Contrastive Predictive Coding," 2018.\n'
    'Banville et al., "Self-supervised representation learning from electroencephalography"\n'
    '  — CPC应用于EEG信号的标杆工作, 证明时序对比预测对生理信号有效\n\n'
    '核心思想: 从过去的ECG token预测未来的PPG token, '
    '并用对比loss区分"真实的未来"和"随机扰动后的未来"。'
    '这和JEPA是互补的——JEPA预测全局嵌入, CPC预测时序变化。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    '新增CPC辅助任务:\n'
    '  1. 取ECG encoder的前半段tokens (t=1~31)\n'
    '  2. 预测PPG后段token的上下文表示 (t=40~50)\n'
    '  3. CPC loss: 真实未来PPG vs 打乱后的伪未来PPG\n'
    '  4. L = L_JEPA + L_InfoNCE + 0.05×L_CPC\n\n'
    '改动量: ~35行'
)

# ══════ 6. Morphology ══════
doc.add_heading('思路六：生理形态学特征约束', level=1)

doc.add_heading('参考论文', level=2)
doc.add_paragraph(
    'NeuroKit2 (Makowski et al., 2021) — 开源生理信号处理库\n'
    'Systole/Cardioception (Legrand et al., 2022) — PPG特征提取\n'
    'CardioPPG (npj Digital Medicine, 2025) — 用形态特征做对比学习\n\n'
    '核心思想: 提取可解释的生理特征(心率、收缩期幅度、反射波指数等), '
    '让编码器的pooled embedding能够预测这些临床特征。'
    'M2AE论文也用了~150维形态特征做正负样本筛选。'
)

doc.add_heading('实现方案', level=2)
doc.add_paragraph(
    '新增形态学预测辅助任务:\n'
    '  1. 用NeuroKit2从PPG/ECG提取20维形态特征\n'
    '    (心率, RR间期std, 收缩期幅度, 反射波延迟, 上升斜率...)\n'
    '  2. 冻结, 从编码器pooled嵌入 → MorphoHead(512→20)\n'
    '  3. L_morph = SmoothL1Loss(预测特征, 真实形态特征)\n'
    '  4. L = L_JEPA + L_InfoNCE + 0.1×L_morph\n\n'
    '改动量: ~30行 (StatsPredHead已存在, 增加特征维度即可)'
)

doc.add_heading('预期收益', level=2)
doc.add_paragraph(
    '让编码器知道自己学到的东西对应什么临床指标; '
    '对CHD分类有直接帮助(反射波指数等是CHD的已知biomarker); '
    '论文中可解释性强'
)

# ══════ Summary ══════
doc.add_heading('七、综合推荐', level=1)

table = doc.add_table(rows=7, cols=6)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['#', '思路', '难度', '改动量', '与当前兼容', '预期收益']):
    table.rows[0].cells[i].text = h
recs = [
    ['1', 'Token级对齐', '低', '~30行', '替换InfoNCE', '消除batch瓶颈'],
    ['2', 'Curriculum Masking', '极低', '~10行', '不改架构', '零风险收益'],
    ['3', '频域掩码', '中', '~40行', '新增辅助任务', 'CHD金标准特征'],
    ['4', '两阶段预训练', '低', '~20行', '加Decoder', '已验证范式'],
    ['5', '时序CPC', '中', '~35行', '新增辅助任务', '时序因果学习'],
    ['6', '形态学约束', '低', '~30行', '已存StatsHead', '可解释性'],
]
for i, row in enumerate(recs):
    for j, v in enumerate(row):
        table.rows[i+1].cells[j].text = v

doc.add_heading('八、8篇参考论文列表', level=1)

papers = [
    ('1', 'Are Electrodes Universal? Resolution-Agnostic Tokenization for Time Series', 'Shi et al.', '2025', 'arXiv', 'Token级跨模态对齐方案'),
    ('2', 'WavesFM: Wavelet-Based Masked Autoencoder for PPG', 'MIT+Harvard', '2025', 'ICLR 2025', '频域+时域双域重建'),
    ('3', 'STH-MAE: Spatial-Temporal Hierarchical Decoupled MAE for ECG', 'ScienceDirect', '2025', 'Expert Systems with Applications', '时空解耦掩码策略'),
    ('4', 'ECGFounder-PT: Post-training Strategy for ECG Foundation Models', 'Li et al.', '2025', 'arXiv', '两阶段训练+Morpholog约束'),
    ('5', 'BioX-Bridge: Unsupervised Cross-Modal Knowledge Transfer across Biosignals', 'Multiple', '2025', 'arXiv', '轻量桥接跨模态学习'),
    ('6', 'Representation Learning with Contrastive Predictive Coding', 'Oord et al.', '2018', 'arXiv经典', '时序对比预测CPC'),
    ('7', 'CardioPPG: AI Modeling PPG to ECG for CVD Prediction', 'npj Digital Medicine', '2025', 'Nature', '形态特征+对比学习'),
    ('8', 'Cross-Modal Representational KD for Spike-Informed LFP', 'Erturk et al.', '2025', 'NeurIPS 2025', 'Token级蒸馏+cosine对齐'),
]
for num, title, authors, year, venue, relevance in papers:
    doc.add_paragraph(
        f'[{num}] {title}\n'
        f'    {authors} | {venue} ({year})\n'
        f'    借鉴点: {relevance}'
    )

output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/Pretrain_Enhancement_Proposals.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
