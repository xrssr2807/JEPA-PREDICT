"""
JEPA-PREDICT 今日修改 + 年龄预测讨论
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

doc.add_heading('JEPA-PREDICT 今日修改与年龄预测讨论', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('2026-06-25').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════ Part 1: Core Problem ═══════════
doc.add_heading('一、核心问题诊断', level=1)

doc.add_paragraph(
    '今日所有下游失败均源于预训练编码器质量不佳，而非下游代码问题。具体原因链：\n\n'
    'batch_size过小 → InfoNCE对比信号不足 → 编码器产出坍塌特征 → 下游全卡死在AUC 0.5\n\n'
    '验证：同一batch=170编码器在train_distill.py(ECG蒸馏提供额外梯度)下可达AUC 0.74，\n'
    '但在标准下游(无蒸馏)下完全崩溃。纯JEPA(JETS+EMA)后期仍会漂移(0.42→0.57)。\n\n'
    '根本解决方案：恢复M2AE原版InfoNCE + 梯度累积，使有效batch=320 > 310(历史最优)。'
)

# ═══════════ Part 2: Architecture Modifications ═══════════
doc.add_heading('二、今日架构修改历程 (4次尝试)', level=1)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['#', '方案', '原理', '结果', '结论']):
    table.rows[0].cells[i].text = h

attempts = [
    ['1', 'VICReg替代InfoNCE\n(VICReg: ICLR 2022)',
     '方差正则(σ≥1) + 协方差去冗余\n不需负样本/大batch',
     'Loss: 1.77→1.48→1.57反弹\n收敛失败',
     '方差阈值对512维过于激进\n破坏JEPA自稳定动态'],
    ['2', '纯JEPA+JETS\n(回退I-JEPA原生)',
     'EMA+stop-gradient天然防坍缩\nJETS拖慢退化',
     'Loss: 0.91→0.42→0.57上升\n后期EMA冻结漂移',
     'EMA单独在生理信号上不够\n需额外约束'],
    ['3', '轻量InfoNCE\n(复用target_embed)',
     'target_embed(256)→投影→128\n无需额外encoder前向',
     'Loss: 6.39→5.62缓慢下降\ntarget侧无梯度',
     '单向对比收敛太慢\n原版M2AE双向前向不可替代'],
    ['4★', 'M2AE+梯度累积\n(最终, 运行中)',
     'restore forward_context(ppg)\nbatch=160, 梯度累积×2\n有效batch=320',
     '运行中 (E0≈5.9)\n预期收敛至0.3~0.5\n未OOM',
     '待验证: eff batch=320\n恢复M2AE对比信号强度'],
]
for i, row in enumerate(attempts):
    for j, v in enumerate(row):
        table.rows[i+1].cells[j].text = v

# ═══════════ Part 3: Code Changes ═══════════
doc.add_heading('三、代码修改明细', level=1)

changes = [
    ('config.py',
     'pretrain_batch_size: 170→310→180→160 (最终值)\n'
     'pretrain_epochs: 50→100\n'
     'use_contrast_loss: False↔True 多次切换, 最终True\n'
     '新增VICReg参数: vicreg_sim/var/cov_weight\n'
     '新增ECG蒸馏参数: use_ecg_distill, distill_lambda\n'
     '新增/移除: downstream_layerdrop (HuBERT-ECG,后废弃)\n'
     '新增: use_dual_channel, use_xgboost'),
    ('models/jepa.py',
     '★ 防坍缩机制多次替换:\n'
     '  1) InfoNCE→VICReg (vicreg_proj_ctx/tgt, 方差+协方差)\n'
     '  2) VICReg→轻量InfoNCE (contrast_proj_ctx/tgt, 复用target_embed)\n'
     '  3) 恢复M2AE InfoNCE (contrast_projector, forward_context(ppg))\n'
     '当前: 原版M2AE, 共享投影头 512→128, InfoNCE τ=0.1'),
    ('models/encoder.py',
     'LayerDrop支持 (HuBERT-ECG风格, 后git checkout恢复)\n'
     '多尺度特征提取 (multiscale_layers=[3,7,11], 后移除)'),
    ('models/classifier.py',
     '新增: DualChannelSimpleFusion (avg_pool→concat→MLP)\n'
     '新增: AsymmetricFusion (ECG冻结+PPG微调, 未使用)\n'
     '新增: unfreeze_ppg_only() 方法\n'
     '新增: MultiScaleClassifier (HiMAE多尺度)'),
    ('models/losses.py',
     'AsymmetricLoss增加one-hot转换 (兼容int label)\n'
     'compute_pos_weight增加多返回值处理'),
    ('train_pretrain.py',
     '★ 梯度累积代码:\n'
     '  accum_steps = 2, loss = loss / accum_steps\n'
     '  每accum_steps步: clip_grad + optimizer.step() + scheduler.step()'),
    ('train_downstream.py',
     '★ ECG蒸馏集成:\n'
     '  加载冻结ECG encoder + 双投影头(512→256→256)\n'
     '  蒸馏模式: uniform LR + 跳过warmup\n'
     '  Probe LR×5 (加速CoT头收敛)\n'
     '  余弦对齐loss: 1-cos(proj_ppg, proj_ecg)\n'
     '新增: F import, EarlyStopping(patience=15)'),
    ('train_distill.py',
     '★ 新增文件: 独立ECG蒸馏微调\n'
     '  ICASSP 2025方案: 投影头+余弦对齐\n'
     '  Probe自适应epoch (维持总步数=5080)\n'
     '  CosineAnnealingLR (无warmup)\n'
     '  batch=128: AUC 0.740 (最佳)'),
    ('train_ensemble.py',
     '★ 新增文件: Ensemble (单通道CoT+双通道Probe加权)\n'
     '  Grid search权重, 最佳w=0.40\n'
     '  AUC 0.768, 不过拟合'),
    ('dataset/data.py',
     'DownstreamDataset新增: 返回UID, SQI门控, 信号对齐\n'
     'DualDownstreamDataset: ECG+PPG配对加载'),
]

for filename, desc in changes:
    doc.add_heading(filename, level=2)
    doc.add_paragraph(desc)

# ═══════════ Part 4: Running Status ═══════════
doc.add_heading('四、当前运行状态', level=1)
doc.add_paragraph(
    '预训练: M2AE + JETS 70% + batch=160 + 梯度累积×2 + 100 epoch\n'
    '  有效batch = 320 (>310历史最优)\n'
    '  预计耗时 ~12h\n'
    '  当前Loss ~5.9 (E0), 预期收敛至0.3~0.5\n\n'
    '下游计划 (预训练完成后):\n'
    '  1. 单通道CoT (标准下游) → 预期AUC ≥0.75\n'
    '  2. ECG蒸馏 (train_distill.py) → 预期AUC ≥0.74\n'
    '  3. Ensemble → 预期AUC ≥0.77'
)

# ═══════════ Part 5: Key Lessons ═══════════
doc.add_heading('五、关键教训', level=1)

lessons = [
    'InfoNCE有效batch是制约因素, 非可选调参——batch<200时对比信号不足以维持收敛',
    'I-JEPA的EMA+stop-gradient在生理信号11.5万对数据上单独不足, 需M2AE辅助',
    'VICReg不适用于JEPA: 方差正则(σ≥1)破坏EMA动态平衡, 导致loss反弹',
    'M2AE使用context encodings(ECG)↔context encodings(PPG)的InfoNCE, 双向前向不可替代',
    '梯度累积是突破显存瓶颈的正确手段, 对训练动态影响极小',
    '双通道全微调(110M params vs 65k samples)必然过拟合, 部署时只用单通道PPG',
    'ECG作为预训练辅助信号是可行且高效的, 双向JEPA反而导致编码器坍塌',
    '所有下游失败排查应从预训练权重的编码器质量开始验证, 而非下游代码',
]
for l in lessons:
    doc.add_paragraph(l, style='List Bullet')

# ═══════════ Part 6: Age Prediction ═══════════
doc.add_heading('六、年龄预测功能扩展', level=1)

doc.add_heading('6.1 架构设计', level=2)
doc.add_paragraph(
    '多任务学习架构 (Multi-Task Learning):\n\n'
    'PPG → Encoder → pool (512)\n'
    '                    ├── CoT Head → CHD分类 (FocalLoss, γ=2)\n'
    '                    └── MLP Head → 年龄回归 (L1 Smooth Loss)\n\n'
    'L_total = L_CHD_cls + λ_age × L_age_reg\n'
    'λ_age建议 = 0.1~0.3 (回归loss量级通常较大, 需缩放)'
)

doc.add_heading('6.2 数据要求', level=2)

table2 = doc.add_table(rows=4, cols=3)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['数据项', '当前状态', '需求']):
    table2.rows[0].cells[i].text = h
age_data = [
    ['PPG信号', '✓ 有 (65k + 11.5万对)', '保持不变'],
    ['CHD标签', '✓ 有 (65k标注)', '保持不变'],
    ['年龄标签', '✗ 无', '每个样本需要一个年龄值 (int/float)'],
]
for i, (k, s, r) in enumerate(age_data):
    table2.rows[i+1].cells[0].text = k
    table2.rows[i+1].cells[1].text = s
    table2.rows[i+1].cells[2].text = r

doc.add_paragraph(
    '\n当前数据文件字段: data, uid, sampling_rate, label — 无年龄信息。\n\n'
    '获取年龄数据的路径:\n'
    '  A. 从uid反向查询原始数据库获取患者元数据(年龄/性别等)\n'
    '  B. 使用公开带年龄的数据集: MIMIC, PPG-DaLiA, PulseDB, WESAD\n'
    '  C. 从预训练数据的文件命名或metadata中提取'
)

doc.add_heading('6.3 收益', level=2)
doc.add_paragraph(
    '多任务学习有助于:\n'
    '  1. 年龄是CHD的强先验(年龄越大风险越高), 共享表征中注入年龄信息可提升CHD诊断\n'
    '  2. "心血管年龄" (模型预测年龄 vs 实际年龄) 本身是心血管健康指标,\n'
    '     差值可作为论文中额外的生物学marker\n'
    '  3. 多任务学习在生理信号文献中是加分项 (e.g. CardioPPG多任务评估)\n'
    '  4. 回归任务提供与分类互补的梯度信号, 有助于防止过拟合'
)

doc.add_heading('6.4 实施优先级', level=2)
doc.add_paragraph(
    '短期:  获取年龄标注 → 加MLP回归头 → 多任务训练 (~1天)\n'
    '中期:  用年龄作为额外预训练目标 (从PPG预测年龄, 增强特征)\n'
    '长期:  多任务扩展到血压/心率/血氧等 (Foundation Model路线)'
)

output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/Today_Report_With_Age_Prediction.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
