"""
JEPA-PREDICT 完整实验总结报告
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

# ── Title ──
title = doc.add_heading('JEPA-PREDICT 完整实验总结报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'JEPA 跨通道 ECG→PPG 预训练 + CHD 冠心病下游分类\n'
    '最终 AUC: 0.768 (Ensemble) / 0.750 (单通道)  |  2026-06-23'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 1. 模型架构 ──
doc.add_heading('1. 模型架构', level=1)

doc.add_heading('1.1 编码器 (SignalEncoder)', level=2)
doc.add_paragraph(
    '输入: (B, 1, L) 单通道生理信号 (L=3000 预训练, L=1000 下游)\n'
    'CNN Stem: 4层1D卷积 [128, 256, 512, 512], 每层 stride=2, 总下采样16×\n'
    'Transformer: 8层, dim=512, 16头, FF=2048, pre-LN, GELU\n'
    '  支持 LayerDrop (随机跳过层, 防过拟合)\n'
    '位置编码: 正弦位置编码 (max_len=200)\n'
    '池化: Adaptive Average Pooling → (B, 512)\n'
    '参数量: ~55M (单编码器) / ~110M (双编码器)'
)

doc.add_heading('1.2 JEPA 预训练架构', level=2)
doc.add_paragraph(
    'Context Encoder (ECG, 可训练) + Target Encoder (PPG, EMA更新, 无梯度)\n'
    'Projection Head: Linear(512→256) + BatchNorm\n'
    'Predictor: 3层MLP, 含 latent variable z~N(0,I), 每样本4次采样取最优\n'
    '预训练任务: ECG → 预测 PPG 的 target embedding (L2 loss)\n'
    'EMA更新: target = momentum × target + (1-momentum) × context\n'
    '\n'
    '★ JETS 掩码策略 (核心创新):\n'
    '  掩码比率: 70% (保留30% patches), Patch大小: 50采样点\n'
    '  作用: 随机掩码时序patch, 强制编码器从局部信息学习全局表征\n'
    '\n'
    '★ M2AE 跨模态对比学习:\n'
    '  类型: 对称 InfoNCE (温度系数 0.1)\n'
    '  正样本对: (ECG_i, PPG_i) 同一文件同时刻采集\n'
    '  负样本对: (ECG_i, PPG_j) 不同文件\n'
    '  L_total = L_JEPA + 0.1 × L_InfoNCE\n'
    '\n'
    '★ 双向 JEPA (最后实验):\n'
    '  Forward: ECG → Predictor → Target(PPG)\n'
    '  Backward: PPG → Predictor → Target(ECG)\n'
    '  L_total = 0.5 × L_fwd + 0.5 × L_bwd + 0.1 × L_contrast\n'
    '  目的: 强制两个编码器的特征空间对齐, 便于下游双通道融合'
)

doc.add_heading('1.3 下游分类架构', level=2)

table = doc.add_table(rows=4, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['架构', '输入', '融合方式', '参数量']):
    table.rows[0].cells[i].text = h
data = [
    ['单通道 CoT', '仅 PPG', '16 reasoning tokens × cross/self-attn', '~55M'],
    ['双通道 SimpleFusion', 'ECG + PPG', 'avg_pool(512+512) → MLP', '~110M'],
    ['双通道 CoT', 'ECG + PPG', 'concat(124 tokens) → CoT', '~110M'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# ── 2. 所有实验 ──
doc.add_heading('2. 完整实验历程', level=1)

doc.add_heading('2.1 预训练实验', level=2)
table = doc.add_table(rows=8, cols=6)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['Run', '方案', 'LR', 'EMA', 'Best Loss', '收敛?']):
    table.rows[0].cells[i].text = h
data = [
    ['1', '纯JEPA', '2.5e-3', '0.9→0.999', '0.493', '✗ E33坍塌到0.73'],
    ['2', '纯JEPA (LR优化)', '5e-4', '0.996→0.999', '0.330', '✗ 后期漂移到0.60'],
    ['3', 'JEPA+Stats+Contrast', '5e-4', '0.996固定', '0.855', '⚠ 辅助loss过重'],
    ['4', 'JETS(70%)+M2AE', '5e-4', '0.996→0.999', '0.184', '✓ 完全收敛, 零退化'],
    ['5', 'JETS+双向JEPA', '5e-4', '0.996→0.999', '0.212', '✓ 完全收敛, 零退化'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

doc.add_heading('2.2 下游微调实验', level=2)
table = doc.add_table(rows=10, cols=5)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['方案', '预训练', 'Probe AUC', 'FT Best AUC', '过拟合?']):
    table.rows[0].cells[i].text = h
data = [
    ['单通道 MLP', 'Run1', '0.645', '0.706', '否'],
    ['单通道 CoT', 'Run4', '0.690', '0.750', '否 ✅'],
    ['双通道 CoT', 'Run4', '—', '0.65', '严重'],
    ['双通道 SimpleFusion', 'Run4', '0.745', '0.776', '是 (E3后下降)'],
    ['双通道 SF + ECG冻结', 'Run4', '—', '0.767', '是'],
    ['双通道 SF + LayerDrop 0.2', 'Run4', '—', '0.771', '是'],
    ['双通道 SF + 双向JEPA', 'Run5', '0.740', '0.769', '是'],
    ['双通道 XGBoost', 'Run4', '—', '0.709', '否 (但AUC低)'],
    ['Ensemble (CoT+SF Probe)', 'Run4', '—', '0.768', '否 ✅'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# ── 3. 关键技术发现 ──
doc.add_heading('3. 关键技术发现', level=1)

findings = [
    ('JETS掩码是稳定性的核心',
     '70%随机掩码强迫编码器从稀疏局部patch推断全局表征, 等效于强正则化。'
     '加入JETS后, 预训练Loss从0.33降至0.184(↓44%), 且E35-E49完全平坦零退化。'
     '这是整个项目中最重要的单一改进。'),
    ('M2AE InfoNCE对比防止表征坍缩',
     '对称InfoNCE在嵌入空间施加结构化约束: 同一人的ECG-PPG必须靠近, '
     '不同人必须推开。提供了纯JEPA缺失的"负样本推力"。'),
    ('EMA固定 vs EMA调度',
     'EMA从0.996→0.999的余弦调度在后期导致target encoder冻结(每步仅0.1%更新), '
     '造成JEPA预测目标漂移。但JETS+M2AE的组合即使使用EMA调度也能保证收敛——'
     '说明JETS掩码是稳定性充分条件。'),
    ('CoT vs MLP 适用场景',
     'CoT头在单通道(62 tokens)下优于MLP(0.72→0.75), 16个reasoning token注意力密度足够。'
     '双通道(124 tokens)时CoT注意力被稀释, 反而MLP向量级融合更有效。'),
    ('双通道必然过拟合——110M参数 vs 65k样本',
     '无论单向/双向JEPA预训练、LayerDrop、冻结部分编码器、XGBoost, '
     '双通道全微调全部过拟合(AUC先升后降)。根本原因: 110M参数/65k样本比 = 1700:1。'
     '双通道收益只能在冻结编码器下拿到(Probe AUC ~0.74-0.75)。'),
    ('双向JEPA未能解决对齐问题',
     '双向预测理论上应强制两个编码器特征空间对齐, 但实测Probe AUC 0.740(与单向0.745相当), '
     'FT过拟合模式完全一致。说明特征质量瓶颈不在空间对齐, 而在数据量。'),
    ('辅助Loss需要极低权重',
     'contrast=0.5 + stats=0.3导致编码器偏向通用特征, CHD特异性下降(Probe AUC 0.62 vs 0.69)。'
     '最优: contrast=0.1, stats关闭。'),
]

for title, text in findings:
    doc.add_heading(title, level=2)
    doc.add_paragraph(text)

# ── 4. 最终结果 ──
doc.add_heading('4. 最终性能指标', level=1)

doc.add_heading('4.1 单通道 CoT (可部署方案)', level=2)
doc.add_paragraph('仅需 PPG 信号, 不过拟合, 训练稳定')

table = doc.add_table(rows=10, cols=2)
table.style = 'Light Grid Accent 1'
for k, v in [
    ('AUC (macro)', '0.7502'), ('Accuracy', '72.28%'),
    ('Precision (macro)', '0.7111'), ('Recall (macro)', '0.6685'),
    ('F1 (macro)', '0.6753'), ('F0.5 (macro)', '0.6924'),
    ('Class 0 Recall (正常)', '87.75%'),
    ('Class 1 Recall (CHD)', '45.96%'),
    ('CHD Precision', '68.81%'),
]:
    table.rows[0].cells[0].text = '指标'
    table.rows[0].cells[1].text = '值'
    break
# rebuild table properly
table2 = doc.add_table(rows=10, cols=2)
table2.style = 'Light Grid Accent 1'
table2.rows[0].cells[0].text = '指标'
table2.rows[0].cells[1].text = '值'
metrics = [
    ('AUC (macro)', '0.7502'), ('Accuracy', '72.28%'),
    ('Precision (macro)', '0.7111'), ('Recall (macro)', '0.6685'),
    ('F1 (macro)', '0.6753'), ('F0.5 (macro)', '0.6924'),
    ('正常 Recall', '87.75%'), ('CHD Recall', '45.96%'),
    ('CHD Precision', '68.81%'),
]
for i, (k, v) in enumerate(metrics):
    table2.rows[i+1].cells[0].text = k
    table2.rows[i+1].cells[1].text = v

doc.add_heading('4.2 Ensemble (最高AUC方案)', level=2)
doc.add_paragraph(
    '单通道CoT(AUC 0.750) + 双通道SimpleFusion Probe(AUC 0.755)\n'
    '加权平均: w_CoT=0.40, w_Dual=0.60\n'
    'AUC: 0.768 | 正常 Recall: 85.2% | CHD Recall: 50.2%'
)

# ── 5. 最优配置 ──
doc.add_heading('5. 最优配置参数', level=1)

table = doc.add_table(rows=12, cols=3)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['参数', '值', '说明']):
    table.rows[0].cells[i].text = h
for i, (p, v, d) in enumerate([
    ('pretrain_lr', '5e-4', '预训练学习率'),
    ('pretrain_epochs', '50', '预训练总epoch'),
    ('pretrain_batch_size', '120-310', '根据显存调整'),
    ('ema_momentum', '0.996→0.999', 'EMA余弦调度'),
    ('jets_mask_ratio', '0.7', 'JETS掩码比率'),
    ('use_contrast_loss', 'True (0.1)', 'M2AE InfoNCE'),
    ('use_cot_head', 'True', '单通道CoT分类头'),
    ('use_layerwise_lr', 'True (0.85)', '逐层LR衰减'),
    ('downstream_epochs', '50-100', '50一般足够'),
    ('downstream_lr', '3e-3', '下游基础LR'),
    ('loss_type', 'focal', 'FocalLoss γ=2'),
]):
    table.rows[i+1].cells[0].text = p
    table.rows[i+1].cells[1].text = v
    table.rows[i+1].cells[2].text = d

# ── 6. 现存问题 ──
doc.add_heading('6. 现存在问题与瓶颈', level=1)

problems = [
    ('下游数据量不足 (核心瓶颈)',
     '65k训练样本对110M双通道模型严重不足(1700:1参数样本比)。双通道全微调必然过拟合, '
     'AUC在2-3个epoch达到峰值后持续下降。所有正则化手段(LayerDrop、冻结编码器、双向预训练)均无法根治。'),
    ('CHD召回率偏低',
     '单通道CoT的CHD召回率仅46%, 超过一半冠心病患者被漏诊。阈值优化后可达80%+, '
     '但以牺牲准确率为代价。双通道Probe在冻结编码器下CHD召回率可达70%+。'),
    ('预训练epoch不足',
     '当前50 epoch, Loss到E49还在微降至0.184(E49=0.185)。翻倍至100 epoch可能进一步提升特征质量, '
     '但时间成本约6-7小时/50 epoch。'),
    ('信号长度不匹配',
     '预训练3000样本 vs 下游1000样本。CNN stride=16导致token数差异(188 vs 62), '
     '位置编码和BatchNorm统计量在两种长度下不一致, 可能降低特征迁移效率。'),
    ('CoT头偶发坍塌',
     'Probe阶段(冻结编码器)CoT偶尔预测恒定输出(P=0.31, R=0.50), '
     '需要较高LR(3e-3)才能稳定训练。'),
    ('双通道评估bug',
     '最终评估有时用错单/双通道接口, 导致best AUC和final AUC不一致(如0.771→0.700)。需修复。'),
    ('SQI质量门控不可用',
     '基于自相关周期性的SQI对1000样本CHD数据输出极低值, 拒绝所有样本。需针对CHD数据重设计。'),
]

for title, text in problems:
    doc.add_heading(title, level=2)
    doc.add_paragraph(text)

# ── 7. 建议优化方向 ──
doc.add_heading('7. 后续优化方向 (按预期收益排序)', level=1)
future = [
    '扩大预训练数据 (最有效): 当前11.5万对, Biosignal Fingerprinting论文340万对。数据量是AUC天花板的主要决定因素。',
    '预训练100 epoch: Loss还在下降趋势中, 翻倍可提升特征质量。预估AUC +0.01-0.02。',
    '下游信号对齐: 将1000样本插值到2000-3000, 匹配预训练尺度, 消除CNN统计量不匹配。',
    '更大规模预训练: 增加公开数据集(MIMIC等), 或使用更多未标注数据。',
    '迁移CardioPPG的PPG→ECG自回归方法: 在PPG-only场景下生成伪ECG特征, 间接利用ECG信息。',
    '多尺度分类头(HiMAE): 不同疾病依赖不同时间尺度特征, 多尺度可捕获更全面信息。',
    '修复并启用双通道XGBoost: 冻结编码器提取特征后, 用更复杂的树模型集成替代简单MLP。',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ── 8. 代码改动清单 ──
doc.add_heading('8. 代码修改清单', level=1)
changes = [
    'config.py: 新增下游参数 (use_cot_head, use_layerwise_lr, layer_decay, '
    'downstream_epochs, downstream_lr, signal_quality_gate, jets_mask_ratio, '
    'jets_mask_patch_size, downstream_layerdrop, use_dual_channel)',
    'config.py: 预训练参数优化 (pretrain_lr→5e-4, warmup→5, ema→0.996, batch→120)',
    'models/encoder.py: 新增 LayerDrop 支持 (TransformerStack + SignalEncoder)',
    'models/encoder.py: 新增多尺度特征提取 (use_multiscale, multiscale_layers)',
    'models/jepa.py: 新增 JETS 掩码策略 (_apply_jets_mask)',
    'models/jepa.py: 新增 M2AE 跨模态对比损失 (_compute_contrast_loss InfoNCE)',
    'models/jepa.py: 新增双向JEPA (ECG→PPG + PPG→ECG 对称预测)',
    'models/classifier.py: 新增 CoT 推理头 (SignalClassifierCoT, LatentReasoningHead)',
    'models/classifier.py: 新增多尺度分类器 (MultiScaleClassifier)',
    'models/classifier.py: 新增双通道分类器 (DualChannelClassifierCoT, DualChannelSimpleFusion, DualChannelClassifier)',
    'models/classifier.py: 新增 XGBoost 支持 (extract_features for XGBoost)',
    'models/losses.py: 新增 AsymmetricLoss, MultiLabelFocalLoss, FocalLoss, compute_pos_weight',
    'train_pretrain.py: 支持 raw/processed 数据切换, 支持 resume',
    'train_downstream.py: 支持双通道训练 (is_dual), SWA, EarlyStopping, TTA, '
    'MixUp, XGBoost路径, Layer-wise LR, 患者级聚合评估, 阈值优化',
    'dataset/data.py: 新增 DualDownstreamDataset, SQI质量门控, UID返回, '
    '信号长度对齐, PhysioAugment',
]

for c in changes:
    doc.add_paragraph(c, style='List Bullet')

# ── Save ──
output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/JEPA_PREDICT_Final_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
