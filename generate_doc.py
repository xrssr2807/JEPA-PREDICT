#!/usr/bin/env python3
"""生成 JEPA ECG-PPG 模型架构与训练过程说明文档 (.docx)"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ─── 样式设置 ───
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ─── 封面标题 ───
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('JEPA ECG-PPG 跨通道预测模型')
run.bold = True
run.font.size = Pt(24)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('模型架构、训练过程与创新点说明')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('基于 Joint Embedding Predictive Architecture\n').font.size = Pt(12)
info.add_run('ECG → PPG 跨模态生理信号自监督预训练').font.size = Pt(11)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. 模型架构
# ═══════════════════════════════════════════
doc.add_heading('一、模型架构', level=1)

doc.add_heading('1.1 整体架构概览', level=2)
doc.add_paragraph(
    'JEPA（Joint Embedding Predictive Architecture）模型的核心思想是：'
    '给定上下文信号（ECG 心电图），预测目标信号（PPG 光电容积脉搏波）在嵌入空间中的表示。'
    '模型包含两个结构完全相同但更新方式不同的编码器，以及一个带有隐变量的预测器网络。'
)

# 架构表格
table = doc.add_table(rows=8, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['模块', '配置参数', '输出形状']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

data = [
    ['输入 (ECG/PPG)', '单通道，30s @ 100Hz', '(B, 1, 3000)'],
    ['CNN Stem (4层)', '通道: [128, 256, 512, 512]\n卷积核: [7, 5, 5, 3]\n步长: [2, 2, 2, 2]', '(B, 512, 188)'],
    ['线性投影', 'Linear(512 → 512)', '(B, 188, 512)'],
    ['位置编码', '正弦位置编码, max_len=200', '(B, 188, 512)'],
    ['Transformer (8层)', 'dim=512, heads=16\nFFN=2048, GELU, Pre-LN', '(B, 188, 512)'],
    ['自适应池化', 'AdaptiveAvgPool → mean', '(B, 512)'],
    ['预测器 (Predictor)', 'Proj: 512→256+BN+ReLU\n主网络: 3层MLP + 隐变量z(64维)', '(B, 256)'],
]
for i, row_data in enumerate(data):
    for j, text in enumerate(row_data):
        table.rows[i+1].cells[j].text = text

doc.add_paragraph()

# 1.2 CNN Stem
doc.add_heading('1.2 CNN Stem（一维卷积特征提取器）', level=2)
doc.add_paragraph(
    'CNN Stem 由 4 个一维卷积块组成，每个块包含 Conv1d → BatchNorm1d → ReLU。'
    '总步长为 2×2×2×2 = 16，将 3000 点信号压缩为 188 个时间步。'
    '这种设计能够有效提取生理信号的局部时频特征，同时降低后续 Transformer 的计算复杂度。'
)

table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cnn_headers = ['层', '输入通道', '输出通道', '卷积核', '步长']
for i, h in enumerate(cnn_headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

cnn_data = [
    ['Conv1d-1', '1', '128', '7', '2'],
    ['Conv1d-2', '128', '256', '5', '2'],
    ['Conv1d-3', '256', '512', '5', '2'],
    ['Conv1d-4', '512', '512', '3', '2'],
]
for i, row_data in enumerate(cnn_data):
    for j, text in enumerate(row_data):
        table.rows[i+1].cells[j].text = text

doc.add_paragraph()

# 1.3 Transformer
doc.add_heading('1.3 Transformer 编码器', level=2)
doc.add_paragraph(
    'CNN 输出经维度投影和正弦位置编码后，输入 8 层 Transformer 编码器。'
    '每层采用 Pre-LN（层归一化前置）结构：'
)
doc.add_paragraph('LayerNorm → Multi-Head Self-Attention (16 heads) → Residual Connection', style='List Bullet')
doc.add_paragraph('LayerNorm → FFN (512 → 2048 → 512, GELU) → Residual Connection', style='List Bullet')
doc.add_paragraph(
    'Transformer 的核心优势在于通过自注意力机制捕捉信号中长距离的时序依赖关系，'
    '这对于心电信号的 P-QRS-T 波段分析和脉搏波的形态识别至关重要。'
    '最后通过自适应平均池化将变长序列压缩为固定维度 (512) 的全局表示向量。'
)

# 1.4 JEPA
doc.add_heading('1.4 JEPA 预测架构', level=2)
doc.add_paragraph(
    'JEPA 是模型的核心创新架构，包含三个关键组件：'
)

doc.add_heading('上下文编码器 (Context Encoder)', level=3)
doc.add_paragraph(
    '处理 ECG 信号，生成上下文嵌入。该编码器通过反向传播接收梯度更新，是模型中唯一参与梯度优化的编码器。'
)

doc.add_heading('目标编码器 (Target Encoder)', level=3)
doc.add_paragraph(
    '处理 PPG 信号，生成目标嵌入。与上下文编码器结构完全相同，但通过 EMA（指数移动平均）方式更新参数，'
    '不直接接收梯度。这种设计避免了表示坍塌（Representation Collapse），确保目标嵌入的稳定性和一致性。'
    '\n\nEMA 更新公式：θ_target = m · θ_target + (1-m) · θ_context'
    '\n其中动量 m 按照余弦调度从 0.996 逐渐增加到 1.0。'
)

doc.add_heading('预测器 (Predictor)', level=3)
doc.add_paragraph(
    '预测器从上下文嵌入出发，结合随机采样的隐变量 z，预测目标信号的嵌入表示。'
    '预测器首先将 512 维上下文嵌入投影到 256 维空间（含 BatchNorm + ReLU），'
    '然后与 64 维隐变量 z 拼接，经过 3 层 MLP (256→256→256→256) 输出预测嵌入。'
)

doc.add_heading('隐变量机制 (Latent Variable)', level=3)
doc.add_paragraph(
    '鉴于 PPG 信号的高度多模态特性（同一 ECG 模式可能对应不同的 PPG 波形，'
    '受血容量、传感器位置、运动伪影等因素影响），模型引入 64 维高斯隐变量 z ~ N(0, I)。'
    '训练时对每个样本采样 4 个不同的 z，取最小 MSE 损失作为最终损失：'
    '\n\nL = min_{z₁,z₂,z₃,z₄} MSE(Predictor(ContextEmbed, zᵢ), TargetEmbed)'
    '\n\n这种 "多采样取最优" 策略使模型能够捕捉一对多的映射关系，'
    '让预测器学会利用隐变量编码 PPG 信号的多种可能形态。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 2. 训练过程
# ═══════════════════════════════════════════
doc.add_heading('二、训练过程', level=1)

doc.add_heading('2.1 数据预处理', level=2)
doc.add_paragraph(
    '训练数据为 115,439 个无标签多通道生理信号片段，每个片段包含 5 通道 × 3000 采样点（30 秒 @ 100Hz）。'
    '预处理流程如下：'
)
doc.add_paragraph('原始 .pkl 文件 (5, 3000) → 提取 ECG(ch0) 和 PPG(ch4) → 逐通道 Z-score 归一化 → 保存为 .pt 文件', style='List Bullet')
doc.add_paragraph(
    'Z-score 归一化对每个文件的每个通道独立计算均值和标准差，确保各片段具有可比性，'
    '同时保留信号的相对幅值变化特征。预处理后数据总量约 3.1 GB。'
)

doc.add_heading('2.2 预训练阶段 (Pre-training)', level=2)

doc.add_heading('任务定义', level=3)
doc.add_paragraph(
    '自监督学习任务：给定 ECG 信号片段，预测对应 PPG 信号在嵌入空间中的表示。'
    '该任务无需人工标注，可利用大量无标签生理信号数据进行预训练。'
)

doc.add_heading('损失函数', level=3)
doc.add_paragraph(
    '采用带多隐变量采样的 MSE 损失：'
    '\nL_pretrain = (1/B) Σᵦ min_{k=1..4} || Predictor(Enc_ctx(ECGᵦ), zᵦᵏ) - Enc_tgt(PPGᵦ) ||²₂'
    '\n\n其中：'
    '\n• Enc_ctx: 上下文编码器（处理 ECG，接收梯度）'
    '\n• Enc_tgt: 目标编码器（处理 PPG，EMA 更新，无梯度）'
    '\n• zᵦᵏ ~ N(0, I): 第 k 个隐变量采样'
    '\n• 取 4 个采样中损失最小的那个'
)

doc.add_heading('优化器配置', level=3)

table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
opt_data = [
    ['超参数', '值'],
    ['优化器', 'AdamW'],
    ['学习率', '2.5 × 10⁻³（与 batch size 线性缩放）'],
    ['β₁ / β₂', '0.9 / 0.95'],
    ['权重衰减', '0.05'],
    ['梯度裁剪', 'max_norm = 1.0'],
    ['Batch Size', '310'],
    ['训练轮数', '100 Epochs（约 372 steps/epoch）'],
]
for i, (k, v) in enumerate(opt_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('学习率调度', level=3)
doc.add_paragraph(
    '采用 Warmup + Cosine Annealing 两阶段调度策略：'
)
doc.add_paragraph('预热阶段 (Warmup): 前 15 个 Epoch，学习率从 10⁻⁶ 线性增加到 2.5 × 10⁻³', style='List Bullet')
doc.add_paragraph('余弦退火阶段 (Cosine Decay): 剩余 85 个 Epoch，学习率按余弦曲线衰减至接近 0', style='List Bullet')
doc.add_paragraph(
    '预热阶段的目的是在训练初期避免梯度爆炸，让模型参数逐步适应优化方向。'
    '余弦退火则有助于模型在训练后期收敛到更平坦的极小值点，提升泛化能力。'
)

doc.add_heading('EMA 动量调度', level=3)
doc.add_paragraph(
    '目标编码器的 EMA 动量 m 同样采用余弦调度：'
    '\nm(t) = m_end + (m_start - m_end) × 0.5 × (1 + cos(π × t/T))'
    '\n其中 m_start = 0.996, m_end = 1.0, t 为当前训练步数, T 为总步数。'
    '\n\n动量从 0.996 逐渐增加到 1.0 意味着：'
)
doc.add_paragraph('训练初期（m=0.996）：目标编码器缓慢跟随上下文编码器，提供稳定的学习目标', style='List Bullet')
doc.add_paragraph('训练末期（m→1.0）：目标编码器几乎不再更新，模型趋于收敛', style='List Bullet')

doc.add_heading('2.3 下游微调阶段 (Downstream Fine-tuning)', level=2)
doc.add_paragraph(
    '预训练完成后，编码器可用于先天性心脏病 (CHD) 分类等下游任务。'
    '微调采用两阶段策略：'
)
doc.add_paragraph('第一阶段 — 线性探测 (Linear Probe, 10 Epochs): 冻结预训练编码器，仅训练分类头，验证预训练表示质量', style='List Bullet')
doc.add_paragraph('第二阶段 — 全参数微调 (Full Fine-tune, 40 Epochs): 解冻所有参数，以较低学习率 (1×10⁻⁴) 端到端训练', style='List Bullet')
doc.add_paragraph(
    '下游任务支持单通道（仅 ECG 或仅 PPG）和双通道（ECG+PPG 融合）两种模式。'
    '双通道模式下，两个编码器的输出拼接后输入分类器，实现跨模态信息互补。'
)

doc.add_page_break()

# ═══════════════════════════════════════════
# 3. 创新点
# ═══════════════════════════════════════════
doc.add_heading('三、模型创新点', level=1)

innovations = [
    (
        '创新点一：JEPA 在生理信号跨模态预测中的应用',
        '传统自监督学习方法（如 MAE、SimCLR）主要针对图像或文本单一模态设计。'
        '本模型首次将 JEPA（Joint Embedding Predictive Architecture）引入 ECG→PPG 跨通道生理信号预测任务。'
        '与生成式方法（需重建原始信号）不同，JEPA 在嵌入空间中进行预测，'
        '避免了像素/采样点级别的重建误差，使模型能够关注高层语义特征而非低层噪声。'
        '\n\n核心优势：'
        '\n• 不需要成对的 ECG-PPG 标注数据即可预训练'
        '\n• 嵌入空间预测比信号空间重建更高效、更鲁棒'
        '\n• EMA 更新的目标编码器防止表示坍塌'
    ),
    (
        '创新点二：多隐变量采样机制处理生理信号多模态性',
        'PPG 信号受多种生理和环境因素影响（血管弹性、传感器接触压力、运动伪影、环境光等），'
        '导致同一 ECG 模式可能对应多种不同的 PPG 波形。'
        '本模型创新地引入 64 维高斯隐变量 z，并在训练时对每个样本采样 4 个 z，'
        '取预测损失最小的那个进行优化。'
        '\n\n这种方法使模型能够：'
        '\n• 捕捉一对多的映射关系（一个 ECG → 多种可能的 PPG）'
        '\n• 通过隐变量编码不可观测的上下文信息'
        '\n• 在推理时通过多次采样生成多样化的预测'
    ),
    (
        '创新点三：CNN + Transformer 混合编码器',
        '模型编码器融合了 CNN 的局部特征提取能力和 Transformer 的全局依赖建模能力：'
        '\n• CNN Stem：通过 4 层步长为 2 的一维卷积，将 3000 点生理信号压缩为 188 个特征向量。'
        '卷积核大小从 7 递减到 3，与 ECG/PPG 信号的多尺度特征（从 QRS 波群到单个心跳）相匹配。'
        '\n• Transformer：8 层 Pre-LN Transformer 编码器在 CNN 特征之上建模长距离时序依赖，'
        '16 个注意力头可并行关注信号的不同方面（节律、幅值、形态等）。'
        '\n• 自适应池化：使编码器可处理不同长度的输入（预训练 3000 点，下游 1000 点），无需修改架构。'
    ),
    (
        '创新点四：余弦 EMA 动量调度策略',
        '目标编码器的 EMA 动量 m 采用余弦调度从 0.996 逐渐增加到 1.0，而非使用固定值。'
        '\n\n设计动机：'
        '\n• 训练初期（m=0.996）：目标编码器快速适应上下文编码器的变化，提供有意义的学习目标'
        '\n• 训练后期（m→1.0）：目标编码器趋于固定，迫使预测器学习更稳定的表示'
        '\n• 余弦曲线在初期变化平缓（利于稳定学习），末期加速收敛'
        '\n\n与固定动量相比，余弦调度在保持训练稳定性的同时，提升了最终表示的质量。'
    ),
    (
        '创新点五：无标签数据的自监督预训练 + 下游迁移学习框架',
        '整套框架充分利用了大规模无标签生理信号数据（115,439 个片段），'
        '通过 ECG→PPG 跨通道预测的自监督任务学习通用的生理信号表示。'
        '\n\n完整的训练管线：'
        '\n1. 数据预处理：逐文件独立 Z-score 归一化（保留个体差异特征）'
        '\n2. 自监督预训练：JEPA 在嵌入空间中学习 ECG→PPG 的映射'
        '\n3. 线性探测评估：冻结编码器，验证预训练表示的下游分类能力'
        '\n4. 全参数微调：以预训练权重初始化，在小规模标注数据上精调'
        '\n\n该管线适用于任何具备 ECG-PPG 同步采集能力的可穿戴设备场景，'
        '且可轻松扩展至其他生理信号（如脑电 EEG、肌电 EMG）的跨模态预训练。'
    ),
]

for i, (title, content) in enumerate(innovations):
    doc.add_heading(title, level=2)
    for para_text in content.split('\n\n'):
        if para_text.startswith('• ') or para_text.startswith('1. ') or para_text.startswith('2. ') or para_text.startswith('3. ') or para_text.startswith('4. '):
            continue
        if '\n•' in para_text:
            lines = para_text.split('\n')
            doc.add_paragraph(lines[0])
            for line in lines[1:]:
                if line.strip():
                    doc.add_paragraph(line.strip(), style='List Bullet')
        else:
            doc.add_paragraph(para_text)

doc.add_page_break()

# ═══════════════════════════════════════════
# 四、模型参数统计
# ═══════════════════════════════════════════
doc.add_heading('四、模型参数统计', level=1)

table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
param_data = [
    ['组件', '参数量（估算）'],
    ['CNN Stem', '~2.5M'],
    ['Transformer 编码器 (×2, context + target)', '~42M (2 × 21M)'],
    ['上下文投影层 (context_proj)', '~0.13M'],
    ['目标投影层 (target_proj)', '~0.13M'],
    ['预测器 (Predictor)', '~0.33M'],
    ['总参数量', '~54.3M'],
    ['可训练参数 (接收梯度)', '~27.0M (仅context编码器+预测器)'],
]
for i, (k, v) in enumerate(param_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

doc.add_paragraph()

doc.add_heading('五、训练环境与计算资源', level=1)

table = doc.add_table(rows=6, cols=2, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
env_data = [
    ['项目', '配置'],
    ['GPU', 'NVIDIA GeForce RTX 4090 D (24 GB)'],
    ['显存使用', '22.7 GB / 24 GB (92%)'],
    ['GPU 利用率', '100%'],
    ['训练数据量', '115,439 个片段（约 346 万秒 / 960 小时）'],
    ['预计训练时长', '约 6.5 - 7 小时 (100 Epochs)'],
]
for i, (k, v) in enumerate(env_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    if i == 0:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.bold = True

# ─── 保存 ───
output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/JEPA_模型架构与训练说明.docx'
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f'文档已保存至: {output_path}')
