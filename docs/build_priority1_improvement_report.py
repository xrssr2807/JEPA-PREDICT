from __future__ import annotations

from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "JEPA-PREDICT模型改进路线与第一优先级实施报告_20260722.docx"
ASSET_DIR = ROOT / "priority1_improvement_assets"
ASSET_DIR.mkdir(exist_ok=True)

NAVY = "17365D"
BLUE = "2E74B5"
TEAL = "287D8E"
GREEN = "4F8A63"
GOLD = "A66A16"
RED = "A33A3A"
INK = "20262E"
MUTED = "66717E"
LIGHT_BLUE = "EAF2F8"
LIGHT_TEAL = "E8F3F4"
LIGHT_GOLD = "FFF5DF"
LIGHT_RED = "FBECEC"
LIGHT_GRAY = "F2F4F7"
WHITE = "FFFFFF"
FONT_CN = "Microsoft YaHei"
FONT_LATIN = "Calibri"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = FONT_LATIN
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    rpr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    rpr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = rgb(color)


def set_paragraph(paragraph, before=0, after=6, line=1.10, keep=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def set_table_geometry(table, widths: list[int], indent=120):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text: str, level: int):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_body(doc, text: str, *, bold_prefix: str | None = None, after=6):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(paragraph.add_run(bold_prefix), bold=True, color=NAVY)
        set_run_font(paragraph.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text: str, level=0):
    paragraph = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    set_paragraph(paragraph, after=4)
    set_run_font(paragraph.add_run(text), size=10.25)
    return paragraph


def add_callout(doc, label: str, text: str, fill=LIGHT_BLUE, accent=BLUE):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=3, after=9, line=1.08)
    paragraph.paragraph_format.left_indent = Inches(0.14)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "5")
    left.set(qn("w:color"), accent)
    borders.append(left)
    ppr.append(borders)
    set_run_font(paragraph.add_run(f"{label}  "), bold=True, color=accent)
    set_run_font(paragraph.add_run(text), size=10.25)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(paragraph, after=0, line=1.0)
        set_run_font(paragraph.add_run(header), size=font_size, bold=True, color=NAVY)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        if row_index % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, "FAFBFC")
        for index, value in enumerate(values):
            paragraph = cells[index].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph(paragraph, after=0, line=1.0)
            set_run_font(paragraph.add_run(str(value)), size=font_size)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_code_block(doc, text: str):
    paragraph = doc.add_paragraph()
    set_paragraph(paragraph, before=2, after=8, line=1.0)
    ppr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F5F7FA")
    ppr.append(shd)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.right_indent = Inches(0.10)
    for line_index, line in enumerate(text.splitlines()):
        if line_index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(8.1)
        run.font.color.rgb = rgb(INK)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(paragraph.add_run("第 "), size=8.5, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)
    set_run_font(paragraph.add_run(" 页"), size=8.5, color=MUTED)


def make_architecture_figure() -> Path:
    path = ASSET_DIR / "priority1_architecture.png"
    color = lambda value: f"#{value}"
    image = Image.new("RGB", (2200, 820), color(WHITE))
    draw = ImageDraw.Draw(image)
    font_regular = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 30)
    font_bold = ImageFont.truetype(r"C:\Windows\Fonts\arialbd.ttf", 32)
    font_small = ImageFont.truetype(r"C:\Windows\Fonts\arial.ttf", 25)

    def box(x1, y1, x2, y2, title, subtitle, fill, outline):
        draw.rounded_rectangle(
            (x1, y1, x2, y2), radius=18,
            fill=color(fill), outline=color(outline), width=5,
        )
        bbox = draw.textbbox((0, 0), title, font=font_bold)
        draw.text(((x1 + x2 - bbox[2] + bbox[0]) / 2, y1 + 28), title, font=font_bold, fill=color(NAVY))
        bbox = draw.textbbox((0, 0), subtitle, font=font_small)
        draw.text(((x1 + x2 - bbox[2] + bbox[0]) / 2, y1 + 86), subtitle, font=font_small, fill=color(INK))

    def arrow(x1, y1, x2, y2, color=BLUE):
        draw.line((x1, y1, x2, y2), fill=f"#{color}", width=7)
        draw.polygon([(x2, y2), (x2 - 22, y2 - 14), (x2 - 22, y2 + 14)], fill=f"#{color}")

    box(80, 300, 380, 510, "Patient bag", "unique segments + mask", LIGHT_BLUE, BLUE)
    box(520, 115, 860, 300, "ECG encoder", "context representation", LIGHT_TEAL, TEAL)
    box(520, 510, 860, 695, "PPG encoder", "target representation", LIGHT_GOLD, GOLD)
    box(1010, 210, 1500, 600, "Disease-conditioned fusion", "per-disease ECG / PPG gate", LIGHT_BLUE, NAVY)
    box(1650, 300, 2080, 510, "Patient MIL", "masked per-disease attention", LIGHT_TEAL, GREEN)
    arrow(380, 405, 505, 210)
    arrow(380, 405, 505, 600)
    arrow(860, 210, 995, 330)
    arrow(860, 600, 995, 485)
    arrow(1500, 405, 1635, 405)
    draw.text((995, 675), "Optional: dual teacher -> PPG student distillation", font=font_regular, fill=color(RED))
    draw.line((1230, 650, 720, 650), fill=color(RED), width=5)
    draw.polygon([(720, 650), (742, 636), (742, 664)], fill=color(RED))
    image.save(path, quality=95)
    return path


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        1: (16, BLUE, 16, 8),
        2: (13, BLUE, 12, 6),
        3: (11.5, NAVY, 8, 4),
    }
    for level, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph(header, after=0)
    set_run_font(header.add_run("JEPA-PREDICT | 模型改进实施报告"), size=8.5, bold=True, color=MUTED)
    add_page_number(section.footer.paragraphs[0])


def build_document():
    doc = Document()
    configure_document(doc)
    figure = make_architecture_figure()

    kicker = doc.add_paragraph()
    set_paragraph(kicker, before=12, after=8)
    set_run_font(kicker.add_run("TECHNICAL IMPLEMENTATION BRIEF"), size=9.5, bold=True, color=TEAL)
    title = doc.add_paragraph()
    set_paragraph(title, after=5)
    set_run_font(title.add_run("JEPA-PREDICT 模型改进路线与\n第一优先级实施报告"), size=24, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    set_paragraph(subtitle, after=14)
    set_run_font(subtitle.add_run("ECG-PPG 生理约束预训练、患者级冠心病识别与 PPG-only 迁移"), size=12.5, color=MUTED)
    for label, value in (
        ("日期", "2026-07-22"),
        ("代码分支", "soft-dtw-token-align"),
        ("本轮状态", "第一优先级已实现并通过本地回归测试"),
    ):
        paragraph = doc.add_paragraph()
        set_paragraph(paragraph, after=2)
        set_run_font(paragraph.add_run(f"{label}："), bold=True, color=NAVY)
        set_run_font(paragraph.add_run(value))

    add_callout(
        doc,
        "核心判断",
        "现阶段不应继续通过堆叠模块追求更低的预训练 Loss。严格划分结果表明，Phase 2/3A 已经证明双通道协同有效，当前主要瓶颈是患者袋重复证据、疾病间共享同一模态融合策略，以及双通道知识无法稳定迁移到 PPG-only。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )

    add_heading(doc, "1. 当前结果与诊断", 1)
    add_table(
        doc,
        ["实验", "通道", "Macro AUC", "CHD AUC", "结论"],
        [
            ["Phase 2 epoch 80", "ECG+PPG", "0.7641", "0.8023", "因果传输对双通道有效"],
            ["Phase 3A task-aware", "ECG+PPG", "0.7891", "0.8083", "CHD 增益约 +0.006"],
            ["Phase 3A task-aware", "PPG-only", "0.7651", "0.7748", "单通道迁移仍是主要短板"],
            ["历史最佳（不同设置）", "ECG+PPG", "0.7848", "0.8440", "划分和 support 不同，不直接比较"],
        ],
        [2350, 1200, 1300, 1300, 3210],
        header_fill=LIGHT_BLUE,
    )
    add_body(doc, "这些结果说明：跨模态 ECG→PPG 学习增强了共享的心脏周期信息，但可能同时压缩了 PPG 独有的血管形态信息。继续增加任务反馈权重未必能解决这一结构性问题。")
    add_bullet(doc, "双通道提升明确，说明 ECG 电活动与 PPG 血流动力学存在可利用的互补信息。")
    add_bullet(doc, "Phase 3A 对 CHD 的严格可比增益较小，暂不建议提高反馈梯度比例。")
    add_bullet(doc, "下一轮预训练创新应围绕 shared/private 解耦和缺失模态鲁棒性，而不是再增加通用 Transformer 层。")

    add_heading(doc, "2. 分阶段改进路线", 1)
    add_table(
        doc,
        ["优先级", "改进内容", "解决问题", "是否需预训练", "当前状态"],
        [
            ["P1", "MIL 掩码、疾病条件化融合、双→单蒸馏、配置清理", "重复证据、共享门控、PPG 迁移", "否", "已实现"],
            ["P2", "Shared-Private JEPA", "保留 PPG 血管形态私有信息", "需要继续或重训", "待实施"],
            ["P3", "模态 dropout 与缺失模态一致性", "部署缺失 ECG 时性能下降", "需要", "待实施"],
            ["P4", "PAT 弱监督与 SQI 软权重", "延迟约束过宽、低质量片段干扰", "需要", "待实施"],
            ["P5", "受控 Phase 3A 任务反馈", "使预训练更贴近 CHD 目标", "建议从 Phase 2 续训", "条件实施"],
        ],
        [950, 2450, 2500, 1650, 1810],
        header_fill=LIGHT_TEAL,
    )
    add_callout(
        doc,
        "实施原则",
        "先用无需重做预训练的下游改动验证收益，再决定是否投入完整预训练。这样可以把结构收益与预训练收益分开归因。",
        fill=LIGHT_TEAL,
        accent=TEAL,
    )

    doc.add_page_break()
    add_heading(doc, "3. 第一优先级：本轮代码修改", 1)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    shape = run.add_picture(str(figure), width=Inches(6.35))
    shape._inline.docPr.set("name", "Priority-one downstream architecture")
    shape._inline.docPr.set(
        "descr",
        "Patient segments with a validity mask pass through separate ECG and PPG "
        "encoders, disease-conditioned modality fusion, and masked patient MIL; "
        "a dual teacher can distill a PPG-only student.",
    )
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(caption, after=10)
    set_run_font(caption.add_run("图 1  第一优先级后的患者级双通道下游架构"), size=9, italic=True, color=MUTED)

    add_heading(doc, "3.1 Patient-MIL 有效片段掩码", 2)
    add_body(doc, "原实现对不足 max_segments 的患者重复采样同一片段，并将重复项视为独立证据。本轮改为每个真实片段只出现一次，其余位置零填充，同时返回布尔 segment_mask。")
    add_bullet(doc, "训练与验证均不再通过重复片段补齐患者袋。")
    add_bullet(doc, "MIL softmax 在无效位置使用掩码，无效位置注意力严格为 0。")
    add_bullet(doc, "所有患者袋至少保留一个有效片段，异常输入会显式报错。")

    add_heading(doc, "3.2 疾病条件化 ECG/PPG 融合", 2)
    add_body(doc, "原双通道头先用一个共享门控融合 ECG 与 PPG，再执行疾病条件化片段注意力，因此九种疾病共享同一种模态融合方式。本轮保留两路表示直到患者级聚合，并为每种疾病独立计算：")
    add_bullet(doc, "ECG 片段注意力与 ECG 患者表示。")
    add_bullet(doc, "PPG 片段注意力与 PPG 患者表示。")
    add_bullet(doc, "基于疾病嵌入的 ECG/PPG 向量门控和跨模态交互项。")
    add_body(doc, "因此心律失常、高血压、冠心病等任务可以学习不同的模态依赖。历史下游教师权重仍可通过 legacy shared-gate 分支加载。")

    add_heading(doc, "3.3 双通道教师蒸馏到 PPG-only", 2)
    add_body(doc, "新增显式的 PPG-only 蒸馏入口。训练数据仍读取 ECG+PPG，学生仅切取 PPG 通道；冻结的双通道教师同时提供疾病 logits 与患者级 embedding。")
    add_code_block(doc, "L = L_supervised\n  + 0.3 * L_binary_logit_distillation\n  + 0.1 * L_patient_embedding_cosine")
    add_body(doc, "蒸馏默认关闭，只有传入 --dual_teacher_checkpoint 时启用，不会改变现有基线命令。温度和两项权重均可通过命令行覆盖。")

    add_heading(doc, "3.4 配置与可复现性", 2)
    add_bullet(doc, "删除 ModelConfig 中重复定义的 use_dual_channel，避免后一个值静默覆盖前一个值。")
    add_bullet(doc, "下游 checkpoint 新增模型架构、教师路径、蒸馏超参数、随机种子和数据划分 provenance。")
    add_bullet(doc, "Phase 3A 反馈训练同步接收 segment_mask，避免后续任务反馈继续使用填充片段。")

    add_heading(doc, "4. 训练方式与重训练要求", 1)
    add_callout(
        doc,
        "是否需要重新预训练",
        "本轮 P1 修改不需要重新进行 JEPA 预训练。可直接使用 Phase 2 或 Phase 3A 预训练权重，重新训练下游分类头和微调编码器。只有后续 Shared-Private JEPA、模态 dropout 和质量感知 transport 才需要继续预训练。",
        fill=LIGHT_BLUE,
        accent=BLUE,
    )
    add_heading(doc, "4.1 训练新的疾病条件化双通道模型", 2)
    add_code_block(doc, "mkdir -p outputs_priority1_dual_seed42\n\npython -u train_downstream.py \\\n+  --checkpoint outputs_phase2_seed42_bs192/jepa_epoch_80.pt \\\n+  --dataset multidisease \\\n+  --multidisease_channel both \\\n+  --multidisease_split splits/multidisease_taskaware_downstream.json \\\n+  --output_dir outputs_priority1_dual_seed42 \\\n+  --mil_batch_size 32 \\\n+  --mil_chunk_size 64 \\\n+  --workers 8 \\\n+  --seed 42 \\\n+  2>&1 | tee outputs_priority1_dual_seed42/downstream_console.log")

    add_heading(doc, "4.2 使用双通道最佳权重蒸馏 PPG-only", 2)
    add_code_block(doc, "mkdir -p outputs_priority1_ppg_distill_seed42\n\npython -u train_downstream.py \\\n+  --checkpoint outputs_phase2_seed42_bs192/jepa_epoch_80.pt \\\n+  --dataset multidisease \\\n+  --multidisease_channel ppg \\\n+  --multidisease_split splits/multidisease_taskaware_downstream.json \\\n+  --dual_teacher_checkpoint outputs_priority1_dual_seed42/downstream_multidisease_best.pt \\\n+  --distill_logit_weight 0.3 \\\n+  --distill_embedding_weight 0.1 \\\n+  --distill_temperature 2.0 \\\n+  --output_dir outputs_priority1_ppg_distill_seed42 \\\n+  --mil_batch_size 32 \\\n+  --mil_chunk_size 64 \\\n+  --workers 8 \\\n+  --seed 42 \\\n+  2>&1 | tee outputs_priority1_ppg_distill_seed42/downstream_console.log")
    add_body(doc, "注意：--dual_teacher_checkpoint 必须是完成下游训练后的双通道 downstream_multidisease_best.pt，而不是 JEPA 预训练权重。")

    doc.add_page_break()
    add_heading(doc, "5. 对照实验与验收标准", 1)
    add_table(
        doc,
        ["实验", "MIL mask", "疾病条件化融合", "双→PPG 蒸馏", "目的"],
        [
            ["E0 原始基线", "否", "否", "否", "锁定历史结果"],
            ["E1 Mask only", "是", "否（legacy）", "否", "验证重复片段偏差"],
            ["E2 新双通道头", "是", "是", "否", "验证疾病特异模态权重"],
            ["E3 PPG baseline", "是", "不适用", "否", "获得公平 PPG 对照"],
            ["E4 PPG distill", "是", "教师使用", "是", "验证双通道知识迁移"],
        ],
        [1500, 1250, 2100, 1550, 2960],
        header_fill=LIGHT_BLUE,
    )
    add_body(doc, "所有实验必须使用相同患者级划分文件、相同预训练权重和相同训练预算，并至少运行 seed=42、3407、2026 三个随机种子。")
    add_bullet(doc, "主要指标：CHD patient-level ROC-AUC，报告三随机种子均值和标准差。")
    add_bullet(doc, "辅助指标：Macro AUC、CHD AUPRC、Precision、Recall、F1、F0.5 与校准误差。")
    add_bullet(doc, "统计报告：患者级 bootstrap 95% 置信区间；测试集只用于最终一次评估。")
    add_callout(
        doc,
        "建议验收线",
        "PPG-only CHD AUC 相对未蒸馏基线提升至少 0.02；新双通道 CHD AUC 相对 Phase 2 严格基线下降不超过 0.005；三个随机种子方向一致。",
        fill=LIGHT_TEAL,
        accent=GREEN,
    )

    add_heading(doc, "6. 后续预训练创新方向", 1)
    add_heading(doc, "6.1 Shared-Private JEPA", 2)
    add_body(doc, "将每种模态拆分为 shared 与 private 表示，仅在 shared 表示上执行 ECG→PPG 因果 transport；PPG-private 分支保留脉搏波上升时间、反射波、血管顺应性和切迹等血管形态信息。可从 Phase 2 最优权重继续训练 20–40 epoch 做可行性验证，论文最终版再从头预训练。")
    add_heading(doc, "6.2 缺失模态与质量感知", 2)
    add_body(doc, "预训练中引入 ECG+PPG、PPG-only、ECG-only 三种输入状态；双通道样本执行 cross-modal transport，单通道样本执行同模态 masked JEPA。SQI 作为连续权重进入 transport loss 和 MIL attention，不再用硬阈值删除样本。")
    add_heading(doc, "6.3 生理延迟弱监督", 2)
    add_body(doc, "由 ECG R 峰与 PPG 波足估计样本级 PAT，为 transport 提供自适应时间窗；高置信心搏使用 Huber 或分布 KL 约束，并单独报告 delay MAE 与相关性。")

    add_heading(doc, "7. 建议论文创新表述", 1)
    add_callout(
        doc,
        "推荐表述",
        "面向冠心病识别的生理约束共享—私有跨模态 JEPA：通过 ECG→PPG 因果延迟传输学习共享电—血流动力学表示，同时保留 PPG 血管形态私有表示，并利用患者级质量感知 MIL 与缺失模态蒸馏支持双通道和 PPG-only 推理。",
        fill=LIGHT_GOLD,
        accent=GOLD,
    )
    add_body(doc, "不建议把“使用 JEPA”“使用 Patient-MIL”或“ECG→PPG 方向性”单独作为创新点，因为已有工作覆盖这些组成部分。论文贡献应集中在因果 transport、shared/private 解耦、质量感知患者聚合及双到单模态迁移的统一设计。")

    add_heading(doc, "8. 参考方向", 1)
    references = [
        "xMAE: Physiology-Aware Masked Cross-Modal Reconstruction for Biosignal Representation Learning. https://arxiv.org/abs/2605.00973",
        "PhysioJEPA: Multimodal high-frequency physiological representation learning. https://openreview.net/forum?id=bdXsfrNaGY",
        "AIM: Self-supervised Learning for Incomplete Multimodal Wearable Sensor Data. https://openreview.net/forum?id=eOATzq7NvI",
        "Promoting cross-modal representations for physiological foundation models. https://openreview.net/forum?id=HNQxrUOvX4",
        "MOMENT: A Family of Open Time-series Foundation Models. https://proceedings.mlr.press/v235/goswami24a.html",
    ]
    for reference in references:
        add_bullet(doc, reference)

    add_heading(doc, "9. 本轮验证记录", 1)
    add_bullet(doc, "python -m py_compile：config.py、dataset/data.py、models/classifier.py、train_downstream.py、train_taskaware_pretrain.py 全部通过。")
    add_bullet(doc, "test_priority1_downstream.py：6 项测试通过，包括掩码不变性、零填充、双通道疾病条件化融合、旧结构兼容和一次蒸馏反向更新。")
    add_bullet(doc, "test_sanity.py：编码器、JEPA、EMA 和端到端训练趋势全部通过。")

    doc.core_properties.title = "JEPA-PREDICT 模型改进路线与第一优先级实施报告"
    doc.core_properties.subject = "ECG-PPG JEPA downstream priority-one improvements"
    doc.core_properties.author = "JEPA-PREDICT Research Team"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
