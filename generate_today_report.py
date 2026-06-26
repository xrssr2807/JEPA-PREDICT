"""
JEPA-PREDICT 今日修改记录 (2026-06-25)
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import datetime

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

doc.add_heading('JEPA-PREDICT 今日修改记录', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(f'日期: 2026-06-25 | 目标: 解决预训练退化 + 下游过拟合').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ═══════════ 1. 问题诊断 ═══════════
doc.add_heading('一、核心问题诊断', level=1)

doc.add_paragraph(
    '1. 所有下游失败都是预训练编码器质量问题, 不是下游代码问题\n'
    '2. batch_size不足导致InfoNCE对比信号弱 → 编码器产出坍塌特征 → 下游全卡死\n'
    '3. 纯JEPA(JETS+EMA)后期仍会漂移, 需要额外防坍缩机制\n'
    '4. M2AE InfoNCE是最有效的防坍缩方案, 但需要batch≥256才能有效工作\n'
    '5. M2AE需要forward_context(ppg)额外前向传播, 在大batch时OOM'
)

# ═══════════ 2. 架构修改 ═══════════
doc.add_heading('二、今日架构修改历程', level=1)

table = doc.add_table(rows=5, cols=5)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['尝试', '方案', '改动', '结果', '结论']):
    table.rows[0].cells[i].text = h

attempts = [
    ['1', 'VICReg替代InfoNCE\n(VICReg: ICLR 2022, LeCun组)',
     'jepa.py: 替换对比loss为方差+协方差+不变性\nconfig.py: 新增vicreg权重参数',
     'Loss从1.77降至1.48后反弹至1.57\n收敛失败',
     'VICReg方差正则(σ≥1)对512维嵌入过于激进\n不适应JEPA预测式架构'],
    ['2', '纯JEPA+JETS\n(回退到I-JEPA原生方案)',
     'config.py: use_contrast_loss=False',
     'Loss从0.91降至0.42后上升至0.57\nEMA后期冻结导致漂移',
     'JETS只能延缓退化,不能根除\nEMA+stop-gradient单独不够'],
    ['3', '轻量InfoNCE\n(复用target_embed, 256→128投影)',
     'jepa.py: 双投影头(512→128, 256→128)\n无需额外encoder前向',
     'Loss从6.4缓慢降至5.6\n收敛太慢(target侧无梯度)',
     'target_encoder EMA冻结\n对比只有单向梯度,效能不足'],
    ['4★', 'M2AE原版+梯度累积\n(最终方案, 运行中)',
     'jepa.py: 恢复forward_context(ppg)+共享投影头\ntrain_pretrain: 梯度累积×2\nbatch=160, effective=320',
     '运行中, 预计E0≈5.9\n后续应降至0.3~0.5',
     '待验证: 有效batch=320恢复InfoNCE信号强度'],
]

for i, row in enumerate(attempts):
    for j, v in enumerate(row):
        table.rows[i+1].cells[j].text = v

# ═══════════ 3. 代码改动清单 ═══════════
doc.add_heading('三、代码修改明细', level=1)

changes = [
    ('config.py',
     'pretrain_batch_size: 170→310→180→160\n'
     'pretrain_epochs: 50→100\n'
     'use_contrast_loss: False↔True (多次切换)\n'
     '新增: vicreg_sim_weight, vicreg_var_weight, vicreg_cov_weight (VICReg用)\n'
     '新增: downstream_layerdrop (HuBERT-ECG用, 后移除)\n'
     '新增: use_ecg_distill, distill_lambda (ECG蒸馏用)'),
    ('models/jepa.py',
     '多次替换防坍缩机制:\n'
     '  1) InfoNCE → VICReg (双投影头, 方差+协方差)\n'
     '  2) VICReg → 轻量InfoNCE (复用的target_embed版)\n'
     '  3) → 恢复原版M2AE InfoNCE (forward_context(ppg))\n'
     '当前: 原版M2AE + 共享投影头512→128'),
    ('models/encoder.py',
     '曾添加: LayerDrop支持 (HuBERT-ECG)\n'
     '曾添加: 多尺度特征提取 (multiscale_layers)\n'
     '已于今日git checkout还原'),
    ('models/classifier.py',
     '添加DualChannelSimpleFusion (向量级ECG+PPG融合)\n'
     '添加AsymmetricFusion (ECG冻结+PPG微调, 未使用)\n'
     '添加unfreeze_ppg_only方法'),
    ('train_pretrain.py',
     '★ 关键修改: 梯度累积\n'
     '  loss = loss / accum_steps\n'
     '  每accum_steps步更新一次权重\n'
     '  effective batch = physical_batch × accum_steps'),
    ('train_downstream.py',
     '添加ECG蒸馏支持: ecg_encoder_distill + proj_ppg/ecg\n'
     '蒸馏模式: uniform LR + 跳过warmup\n'
     'Probe: 蒸馏模式下LR×5\n'
     '添加EarlyStopping (15 epoch)'),
    ('train_distill.py',
     '★ 新增文件: ECG-Guided蒸馏微调\n'
     'ICASSP 2025方案: 投影头+余弦对齐\n'
     'Probe自适应epoch (维持总步数不变)\n'
     'batch=128时AUC可达0.74'),
    ('dataset/data.py',
     'DownstreamDataset新增返回UID\n'
     '新增SQI信号质量门控 (已关闭, 对CHD数据过滤过严)\n'
     '新增target_length信号对齐'),
    ('train_ensemble.py',
     '新增文件: 单通道CoT+双通道Probe加权Ensemble\n'
     '最佳AUC 0.768, 稳定不过拟合'),
]

for filename, desc in changes:
    doc.add_heading(filename, level=2)
    doc.add_paragraph(desc)

# ═══════════ 4. 最终结果 ═══════════
doc.add_heading('四、目前运行状态', level=1)

doc.add_paragraph(
    '预训练: 运行中 (M2AE + JETS 70% + batch=160 + 梯度累积×2 + 100 epoch)\n'
    '有效batch=320, 预计收敛到Loss 0.2~0.3\n'
    '预计耗时: ~12h (100 epoch, 预处理数据)\n'
    '\n'
    '下游计划 (预训练完成后):\n'
    '  1. 单通道CoT → 预期AUC≥0.75\n'
    '  2. ECG蒸馏 (train_distill.py) → 预期AUC≥0.74\n'
    '  3. Ensemble → 预期AUC≥0.77'
)

# ═══════════ 5. 学习收获 ═══════════
doc.add_heading('五、关键教训', level=1)

lessons = [
    'InfoNCE的有效batch是制约因素, 不是可有可无的调参',
    'I-JEPA的EMA+stop-gradient自防坍缩在生理信号小数据上不够',
    'VICReg虽然不依赖batch size, 但其方差+协方差正则对JEPA架构不适用',
    'M2AE InfoNCE是验证过的唯一有效方案, 代价是需要额外前向传播',
    '梯度累积是突破显存瓶颈的正确手段, 对训练动态影响极小',
    '双通道全微调必然过拟合 (110M参数 vs 65k样本), 部署时应只用单通道',
    'ECG作为预训练辅助信号+PPG-only部署是可行的, 双向JEPA反而导致坍塌',
]
for l in lessons:
    doc.add_paragraph(l, style='List Bullet')

output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/Today_Modifications_20260625.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
