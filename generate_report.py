"""
Generate summary report of JEPA-PREDICT experiments as .docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

# ── Title ──
title = doc.add_heading('JEPA-PREDICT 项目总结报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'JEPA (Joint Embedding Predictive Architecture) 跨通道 ECG→PPG 预训练'
    ' + CHD 冠心病下游分类',
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    f'最终 AUC: 0.7502 | 日期: 2026-06-22'
).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── 1. Architecture ──
doc.add_heading('1. 模型架构', level=1)

doc.add_heading('1.1 编码器 (SignalEncoder)', level=2)
doc.add_paragraph(
    '输入: (B, 1, L) 单通道生理信号 (L=3000 预训练, L=1000 下游)\n'
    'CNN Stem: 4层1D卷积 [128, 256, 512, 512], stride=2/层, 总下采样16×\n'
    'Transformer: 8层, dim=512, 16头, FF=2048, pre-LN, GELU\n'
    '位置编码: 正弦位置编码 (max_len=200)\n'
    '池化: Adaptive Average Pooling → (B, 512)\n'
    '参数量: ~55M'
)

doc.add_heading('1.2 JEPA 预训练架构', level=2)
doc.add_paragraph(
    'Context Encoder (ECG, 可训练) + Target Encoder (PPG, EMA更新, 无梯度)\n'
    'Projection Head: Linear(512→256) + BatchNorm\n'
    'Predictor: 3层MLP [256+64→256→256→256], 含latent variable z~N(0,I)\n'
    '  每样本采样4次z, 取最优 (最小化预测误差)\n'
    '预训练任务: ECG → 预测 PPG 的 target embedding (L2 loss)\n'
    'EMA更新: target = momentum × target + (1-momentum) × context'
)

doc.add_heading('1.3 JETS 掩码策略', level=2)
doc.add_paragraph(
    '掩码比率: 70% (保留30% patches)\n'
    'Patch大小: 50采样点 (3000/50=60 patches)\n'
    '作用: 随机掩码时序patch, 强制编码器从局部信息学习全局表征\n'
    '效果: 预训练Loss从0.33降至0.184, 完全消除后期退化'
)

doc.add_heading('1.4 M2AE 跨模态对比学习', level=2)
doc.add_paragraph(
    '类型: 对称InfoNCE (温度系数0.1)\n'
    '正样本对: (ECG_i, PPG_i) 同一文件同时刻采集\n'
    '负样本对: (ECG_i, PPG_j) 不同文件\n'
    '投影头: Linear(512→512)→BN→GELU→Linear(512→128)\n'
    'L_total = L_JEPA + 0.1 × L_InfoNCE\n'
    '作用: 拉近同一样本的ECG-PPG表征, 推远不同样本, 防止表征坍缩'
)

doc.add_heading('1.5 下游分类头 (CoT)', level=2)
doc.add_paragraph(
    'Chain-of-Thought (CoT) 推理分类头:\n'
    '16个可学习 Reasoning Tokens × 512 dim\n'
    'Cross-Attention: Reasoning Tokens ← Encoder Tokens\n'
    'Self-Attention: Tokens之间交互推理\n'
    'Pooling Query → Decision Token → Linear → (B, 2)\n'
    'Layer-wise LR Decay: 顶层 lr=3e-4, CNN stem lr=4.3e-5 (decay=0.85)'
)

# ── 2. 实验历程 ──
doc.add_heading('2. 实验历程与关键发现', level=1)

table = doc.add_table(rows=8, cols=7)
table.style = 'Light Grid Accent 1'
headers = ['Run', '预训练方案', '下游方案', '预训练Best Loss', '下游AUC', '收敛?', '关键问题']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h

data = [
    ['1', '纯JEPA\nEMA 0.9→0.999\nLR=2.5e-3', 'MLP\nFocalLoss\n50 epoch', '0.493', '0.706', '✗', 'Loss震荡\nE33坍塌至0.73'],
    ['2', '纯JEPA\nEMA 0.996→0.999\nLR=5e-4', '—', '0.330', '—', '✗', '后期漂移\n0.33→0.60'],
    ['3', '同Run2权重', 'CoT+LayerLR\nFocalLoss\n100 epoch', '—', '0.720', '—', 'CoT偶发坍塌\nAUC仍上升'],
    ['4', 'JEPA+aux losses\nEMA=0.996固定', '—', '0.855', '—', '⚠', 'aux loss过重\n特征CHD特异性下降'],
    ['5', '纯JEPA\n(同Run2)', 'CoT+LayerLR\n100 epoch', '0.336', '0.734', '✗', 'Loss先降后升\n中期自行恢复'],
    ['6', 'JETS(70%)+M2AE\nEMA 0.996→0.999\nLR=5e-4', 'CoT+LayerLR\nFocalLoss\n100 epoch\nSQI关闭', '0.184', '0.750', '✓', '首此真正收敛\nE35-E49平坦'],
    ['最佳', 'JETS+M2AE✅', 'CoT+LayerLR ✅', '0.184 ✅', '0.750 ✅', '✓', '—'],
]
for i, row in enumerate(data):
    for j, val in enumerate(row):
        table.rows[i+1].cells[j].text = val

# ── 3. Key Findings ──
doc.add_heading('3. 关键技术发现', level=1)

findings = [
    ('EMA动量调度',
     'EMA从0.996→0.999的余弦调度在后期会导致target encoder冻结(每步仅0.1%更新), '
     '造成JEPA预测目标漂移和Loss退化。固定EMA=0.996(每步0.4%更新)可消除退化但收敛变慢。'
     'JETS+M2AE的组合即使使用EMA调度也能保证收敛。'),
    ('JETS掩码是稳定性的核心',
     '70%随机掩码强迫编码器从稀疏的局部patch推断全局表征, 等效于强正则化, '
     '防止编码器走捷径。预训练Loss从0.33→0.184(↓44%), 且E35→E49完全平坦无退化。'),
    ('M2AE InfoNCE对比防止坍缩',
     '对称InfoNCE在嵌入空间施加结构化约束: 同一人的ECG-PPG必须靠近, '
     '不同人的必须推开。这提供了纯JEPA缺失的"负样本推力", '
     '是表征空间不坍缩的关键保障。'),
    ('SQI质量门控对CHD数据不适用',
     '基于自相关周期性的SQI评分对1000样本的短PPG信号输出极低值(0.05-0.15), '
     '导致所有CHD样本被过滤。CHD患者本身的病理波形即被误判为"低质量"。'
     '关闭SQI后训练恢复正常。'),
    ('CoT推理头比MLP头有显著增益',
     'CoT头(16 reasoning tokens + cross/self-attention) vs MLP头([512→128→64→2]): '
     'AUC提升约+0.03。但CoT在Probe阶段(冻结编码器)易坍塌, 需要足够高的LR。'),
    ('辅助Loss权重需要精细调节',
     'contrast=0.5 + stats=0.3过重导致编码器偏向通用特征, CHD特异性下降(Probe AUC 0.62 vs 0.69)。'
     '当前最优: contrast=0.1(仅M2AE风格InfoNCE), stats关闭。'),
]

for title, text in findings:
    doc.add_heading(title, level=2)
    doc.add_paragraph(text)

# ── 4. Final Metrics ──
doc.add_heading('4. 最终性能指标', level=1)

metrics_table = doc.add_table(rows=11, cols=2)
metrics_table.style = 'Light Grid Accent 1'
metrics_data = [
    ('AUC (macro)', '0.7502'),
    ('Accuracy', '72.28%'),
    ('Precision (macro)', '0.7111'),
    ('Recall (macro)', '0.6685'),
    ('F1 (macro)', '0.6753'),
    ('F0.5 (macro)', '0.6924'),
    ('Class 0 (正常) Precision', '0.7341'),
    ('Class 0 (正常) Recall', '0.8775'),
    ('Class 1 (CHD) Precision', '0.6881'),
    ('Class 1 (CHD) Recall', '0.4596'),
]
for i, (k, v) in enumerate(metrics_data):
    metrics_table.rows[i].cells[0].text = k
    metrics_table.rows[i].cells[1].text = v

# ── 5. Config ──
doc.add_heading('5. 最优配置参数', level=1)

config_table = doc.add_table(rows=12, cols=3)
config_table.style = 'Light Grid Accent 1'
config_headers = ['参数', '值', '说明']
for i, h in enumerate(config_headers):
    config_table.rows[0].cells[i].text = h

config_data = [
    ('pretrain_lr', '5e-4', '预训练学习率'),
    ('pretrain_warmup_epochs', '5', '预热epoch数'),
    ('pretrain_epochs', '50', '预训练总epoch'),
    ('pretrain_batch_size', '170', '预训练batch size (JETS+M2AE需降低)'),
    ('ema_momentum', '0.996', 'EMA初始动量'),
    ('ema_end_momentum', '0.999', 'EMA最终动量'),
    ('jets_mask_ratio', '0.7', 'JETS掩码比率'),
    ('use_contrast_loss', 'True (0.1)', 'M2AE InfoNCE对比'),
    ('use_cot_head', 'True', 'CoT推理分类头'),
    ('use_layerwise_lr', 'True (decay=0.85)', '逐层学习率衰减'),
    ('downstream_epochs', '100', '下游总epoch (10 probe + 90 FT)'),
]
for i, (p, v, d) in enumerate(config_data):
    config_table.rows[i+1].cells[0].text = p
    config_table.rows[i+1].cells[1].text = v
    config_table.rows[i+1].cells[2].text = d

# ── 6. Future ──
doc.add_heading('6. 后续优化方向', level=1)
future = [
    '预训练100 epoch (当前仅50, Loss还在下降趋势中)',
    '修复CoT偶发坍塌: 增大Probe阶段LR或改用渐进解冻',
    '多尺度分类头 (HiMAE MultiScaleClassifier, 已实现)',
    '信号对齐: 下游插值至3000匹配预训练尺度',
    '多通道融合: ECG+PPG DualChannelClassifier (需匹配的ECG数据)',
    'XGBoost替代微调 (已实现, use_xgboost=True)',
    'Ensemble多checkpoint平均预测',
    '预估天花板: AUC ~0.77 (单通道PPG信息量上限)',
]
for f in future:
    doc.add_paragraph(f, style='List Bullet')

# ── Save ──
output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/JEPA_PREDICT_Report.docx'
doc.save(output_path)
print(f'Report saved to {output_path}')
