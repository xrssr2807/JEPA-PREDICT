"""
JEPA-PREDICT 论文计划与进度报告
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.size = Pt(11)
style.font.name = 'Arial'

# ── Title ──
doc.add_heading('JEPA-PREDICT 论文计划与实验进度报告', level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('2026-06-25').alignment = WD_ALIGN_PARAGRAPH.CENTER

# ══════════════════════════════════════════════
# Part 1: Progress Summary
# ══════════════════════════════════════════════
doc.add_heading('一、当前实验进度', level=1)

doc.add_heading('1.1 已完成工作', level=2)

table = doc.add_table(rows=10, cols=4)
table.style = 'Light Grid Accent 1'
for i, h in enumerate(['模块', '内容', '完成度', '关键结果']):
    table.rows[0].cells[i].text = h
for row in table.rows[1:]:
    for c in row.cells: c.text = ''
completed = [
    ['预训练', '5轮消融实验\n(纯JEPA→JETS+M2AE)', '100%',
     'JETS掩码是收敛充要条件\nBest Loss: 0.184\n零退化,完全收敛'],
    ['下游分类', '9种下游方案对比\n(MLP/CoT/双通道/Ensemble等)', '100%',
     '最佳AUC: 0.768 (Ensemble)\n单通道AUC: 0.750'],
    ['消融实验', 'EMA调度/JETS掩码/M2AE对比\n/双向JEPA/辅助Loss', '100%',
     'JETS+对比=收敛充分条件\n双向JEPA=编码器坍塌\n辅助Loss过重=特征退化'],
    ['代码框架', '完整训练pipeline\n+评估+可视化', '100%',
     '支持预训练/微调/XGBoost/\n蒸馏/Ensemble等多种模式'],
]
for i, (m, c, p, r) in enumerate(completed):
    table.rows[i+1].cells[0].text = m
    table.rows[i+1].cells[1].text = c
    table.rows[i+1].cells[2].text = p
    table.rows[i+1].cells[3].text = r

doc.add_heading('1.2 最终性能', level=2)
table2 = doc.add_table(rows=5, cols=4)
table2.style = 'Light Grid Accent 1'
for i, h in enumerate(['方案', 'AUC', 'CHD召回率', '部署要求']):
    table2.rows[0].cells[i].text = h
results = [
    ['单通道 CoT', '0.750', '46%', '仅需PPG ✅'],
    ['双通道 Ensemble', '0.768', '50%', '需ECG+PPG'],
    ['单通道 CoT(双向预训练)', '0.773*', '—', '编码器坍塌,不可用'],
    ['ECG蒸馏(当前)', '0.720', '—', '仅需PPG,但编码器弱'],
]
for i, row in enumerate(results):
    for j, v in enumerate(row):
        table2.rows[i+1].cells[j].text = v

doc.add_heading('1.3 待完成工作', level=2)
pending = [
    '重跑预训练 (batch≥180, 编码器质量验证) — 当前batch=170编码器偏弱',
    'ECG蒸馏微调在好编码器上复现 — 预期AUC 0.76+',
    '与Baseline方法对比 (SVM/XGBoost/纯监督CNN/SimCLR)',
    '外部数据集验证 (MIMIC或PTB-XL公开数据)',
    '论文撰写',
]
for p in pending:
    doc.add_paragraph(p, style='List Bullet')

# ══════════════════════════════════════════════
# Part 2: Paper Plan
# ══════════════════════════════════════════════
doc.add_heading('二、论文计划', level=1)

doc.add_heading('2.1 论文标题 (建议)', level=2)
doc.add_paragraph(
    '"JETS-JEPA: Masked Bidirectional Cross-Modal Self-Supervised Learning '
    'for Collapse-Free PPG-Based Cardiovascular Screening"'
)

doc.add_heading('2.2 核心贡献 (3点)', level=2)
contributions = [
    ('贡献1 — JETS掩码适配1D生理信号',
     '首次将图像域JETS掩码(70%随机patch丢弃)适配到1D ECG/PPG信号，'
     '并通过5轮消融实验证明：掩码是预训练收敛的充要条件。'
     '纯JEPA预训练普遍存在Loss退化(0.33→0.60)，加入JETS后首次实现零退化收敛(0.184)。'),
    ('贡献2 — 跨模态M2AE对比防坍缩',
     'ECG↔PPG的对称InfoNCE对比学习提供"负样本推力"，'
     '与JETS掩码协同工作：掩码阻断捷径，对比防止表征坍缩。'
     '消融实验证明两者单独使用均不足以保证收敛。'),
    ('贡献3 — 小数据场景下的稳定预训练',
     '在仅11.5万对(远小于Biosignal Fingerprinting的340万对)数据上，'
     '通过JETS+M2AE协同实现稳定收敛。预训练Loss 0.184，'
     '下游CHD筛查AUC 0.768(Ensemble)/0.750(单通道PPG-only)。'),
]
for title, text in contributions:
    doc.add_heading(title, level=3)
    doc.add_paragraph(text)

doc.add_heading('2.3 目标期刊', level=2)
journals = [
    'IEEE Journal of Biomedical and Health Informatics (JBHI) — SCI二区, IF~5.0',
    'Biomedical Signal Processing and Control (BSPC) — SCI二区, IF~5.1',
    'Computers in Biology and Medicine (CBM) — SCI二区, IF~4.7',
    '备选: Physiological Measurement, Frontiers in Cardiovascular Medicine',
]
for j in journals:
    doc.add_paragraph(j, style='List Bullet')

doc.add_heading('2.4 论文结构', level=2)
sections = [
    ('1. Introduction',
     '心血管疾病筛查需求 → 可穿戴PPG的潜力与局限 → ECG辅助PPG的思路'),
    ('2. Related Work',
     'CardioPPG, Biosignal Fingerprinting, HuBERT-ECG, JEPA系列'),
    ('3. Method',
     '3.1 JETS Masking for 1D Signals\n'
     '3.2 M2AE Cross-Modal Contrastive Learning\n'
     '3.3 Bidirectional JEPA Architecture\n'
     '3.4 Downstream Fine-tuning Pipeline'),
    ('4. Experiments',
     '4.1 Dataset (11.5万对预训练 + 6.5万CHD标注)\n'
     '4.2 Implementation Details\n'
     '4.3 Pretraining Convergence Analysis (消融实验核心)\n'
     '4.4 Downstream CHD Classification\n'
     '4.5 Ablation Study (EMA/JETS/Contrast/双向)\n'
     '4.6 Comparison with Baselines (待完成)'),
    ('5. Discussion',
     '为何JETS+对比协同有效 / 小数据预训练的启示 / 局限性'),
    ('6. Conclusion', '总结三个贡献 + 未来方向'),
]
for title, desc in sections:
    doc.add_heading(title, level=3)
    doc.add_paragraph(desc)

doc.add_heading('2.5 实验计划 (Table/Figure 分配)', level=2)
figs = [
    'Fig 1: JETS-JEPA 架构图 (预训练+下游全流程)',
    'Fig 2: 预训练Loss曲线对比 (5轮实验的消融, 这是最核心的图)',
    'Table 1: 下游CHD筛查结果 (9种方案对比)',
    'Table 2: 消融实验 (JETS on/off, M2AE on/off, 双向 on/off)',
    'Fig 3: 不同标注数据量下的AUC (标签效率, 对标CardioPPG)',
    'Table 3: 与Baseline方法对比 (待完成)',
    'Fig 4: t-SNE特征可视化 (预训练前后对比)',
]
for f in figs:
    doc.add_paragraph(f, style='List Bullet')

# ══════════════════════════════════════════════
# Part 3: Quick Start
# ══════════════════════════════════════════════
doc.add_heading('三、重现最佳结果的代码指令', level=1)

doc.add_paragraph(
    '# Step 1: 预训练 (~3.5h)\n'
    'python train_pretrain.py\n'
    '# 配置: pretrain_lr=5e-4, warmup=5, EMA=0.996, batch=180+\n'
    '# 输出: outputs/jepa_best.pt\n\n'
    '# Step 2: 单通道CoT下游 (~1h)\n'
    'python train_downstream.py --checkpoint outputs/jepa_best.pt --dataset chd\n'
    '# 配置: use_cot_head=True, use_layerwise_lr=True, use_dual_channel=False\n'
    '# 输出: outputs/downstream_chd_best.pt (AUC ~0.75)\n\n'
    '# Step 3: Ensemble (~30min)\n'
    'python train_ensemble.py\n'
    '# 双通道SimpleFusion Probe + 单通道CoT加权平均\n'
    '# 输出: AUC ~0.768\n\n'
    '# Step 4 (可选): ECG蒸馏微调\n'
    'python train_distill.py\n'
    '# 训练时ECG引导, 部署时仅PPG, 适合可穿戴场景'
)

# ══════════════════════════════════════════════
# Part 4: Key Config
# ══════════════════════════════════════════════
doc.add_heading('四、关键配置参数', level=1)

table3 = doc.add_table(rows=12, cols=3)
table3.style = 'Light Grid Accent 1'
for i, h in enumerate(['参数', '最佳值', '说明']):
    table3.rows[0].cells[i].text = h
configs = [
    ('pretrain_lr', '5e-4', '预训练学习率'),
    ('jets_mask_ratio', '0.7', 'JETS掩码比例 (核心参数)'),
    ('use_contrast_loss', 'True (0.1)', 'M2AE InfoNCE对比'),
    ('ema_momentum', '0.996→0.999', 'EMA余弦调度'),
    ('use_cot_head', 'True', 'CoT推理分类头'),
    ('use_layerwise_lr', 'True (0.85)', '逐层LR衰减'),
    ('downstream_epochs', '100', 'Probe 10 + FT 90'),
    ('downstream_lr', '3e-3', '下游基础LR'),
    ('loss_type', 'focal', 'FocalLoss γ=2'),
    ('use_ecg_distill', 'True (0.3)', 'ECG蒸馏 (实验性)'),
    ('pretrain_batch_size', '180-310', '越大越好, 受显存限制'),
]
for i, (p, v, d) in enumerate(configs):
    table3.rows[i+1].cells[0].text = p
    table3.rows[i+1].cells[1].text = v
    table3.rows[i+1].cells[2].text = d

# ── Save ──
output_path = '/root/autodl-tmp/JEPA-PREDICT/outputs/Paper_Plan_Progress_Report.docx'
doc.save(output_path)
print(f'Saved → {output_path}')
